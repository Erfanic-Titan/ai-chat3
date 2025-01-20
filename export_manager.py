# export_manager.py
import os
from typing import List, Dict
from datetime import datetime
from fpdf import FPDF
import docx
from docx.shared import Pt
import markdown
import logging

logger = logging.getLogger(__name__)

class ExportManager:
    def __init__(self):
        """Initialize export manager"""
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)
        
    def _format_messages(self, messages: List[Dict], format_type: str) -> str:
        """Format messages based on export type"""
        formatted_text = ""
        
        if format_type in ['txt', 'md']:
            for msg in messages:
                role_icon = "👤" if msg["role"] == "user" else "🤖"
                timestamp = datetime.fromisoformat(msg["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
                
                if format_type == 'md':
                    formatted_text += f"### {role_icon} {msg['role'].title()} - {timestamp}\n\n"
                    formatted_text += f"{msg['content']}\n\n---\n\n"
                else:
                    formatted_text += f"{role_icon} {msg['role'].upper()} - {timestamp}\n"
                    formatted_text += f"{msg['content']}\n"
                    formatted_text += "-" * 80 + "\n\n"
                    
        return formatted_text
    
    async def export_chat(
        self,
        chat_id: int,
        messages: List[Dict],
        format_type: str,
        chat_title: str
    ) -> str:
        """Export chat to specified format"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{chat_title}_{timestamp}"
            
            if format_type == "pdf":
                return await self._export_to_pdf(filename, messages)
            elif format_type == "docx":
                return await self._export_to_docx(filename, messages)
            elif format_type == "txt":
                return await self._export_to_txt(filename, messages)
            elif format_type == "md":
                return await self._export_to_markdown(filename, messages)
            else:
                raise ValueError(f"Unsupported format: {format_type}")
                
        except Exception as e:
            logger.error(f"Error exporting chat: {str(e)}")
            raise
    
    async def _export_to_pdf(self, filename: str, messages: List[Dict]) -> str:
        """Export chat to PDF with robust space checking"""
        output_path = os.path.join(self.export_dir, f"{filename}.pdf")
        
        try:
            # Initialize PDF with A4 format and margins
            pdf = FPDF(format='A4')
            
            # Add first page
            pdf.add_page()
            
            # Set initial margins (all 20 mm)
            margin = 20
            pdf.set_margins(margin, margin, margin)
            
            # Calculate effective page width
            effective_width = pdf.w - (2 * margin)
            
            # Cover page
            pdf.set_font("Helvetica", "B", 20)
            pdf.cell(effective_width, 20, "Chat Export", ln=True, align='C')
            
            pdf.ln(10)
            pdf.set_font("Helvetica", "", 12)
            pdf.cell(effective_width, 10, 
                    f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                    ln=True, align='C')
            
            pdf.ln(10)
            pdf.cell(effective_width, 10, 
                    f"Total Messages: {len(messages)}", 
                    ln=True, align='C')
            
            # Content pages
            pdf.add_page()
            pdf.set_font("Helvetica", "", 10)
            
            def write_line_safely(text: str, height: int = 5):
                """Write a line of text safely, breaking it if necessary"""
                if pdf.get_string_width(text) <= effective_width:
                    # Text fits in one line
                    pdf.cell(effective_width, height, text, ln=True)
                else:
                    # Need to break text into smaller chunks
                    words = text.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " + word if current_line else word)
                        
                        if pdf.get_string_width(test_line) <= effective_width:
                            current_line = test_line
                        else:
                            if current_line:
                                pdf.cell(effective_width, height, current_line, ln=True)
                            current_line = word
                    
                    if current_line:  # Write any remaining text
                        pdf.cell(effective_width, height, current_line, ln=True)
            
            # Process each message
            for msg in messages:
                # Check page space
                if pdf.get_y() > pdf.h - 40:
                    pdf.add_page()
                
                # Write header
                pdf.set_font("Helvetica", "B", 10)
                role = "User" if msg["role"] == "user" else "Assistant"
                timestamp = datetime.fromisoformat(msg["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
                header_text = f"{role} - {timestamp}"
                write_line_safely(header_text)
                
                pdf.ln(2)
                
                # Write content
                pdf.set_font("Helvetica", "", 10)
                
                # Split content into lines
                content_lines = msg["content"].split("\n")
                for line in content_lines:
                    if line.strip():  # Only process non-empty lines
                        write_line_safely(line.strip())
                    else:
                        pdf.ln(3)  # Empty line spacing
                
                # Add separator
                pdf.ln(3)
                pdf.line(margin, pdf.get_y(), pdf.w - margin, pdf.get_y())
                pdf.ln(5)
            
            # Add page numbers
            page_count = pdf.page_no()
            for page in range(1, page_count + 1):
                pdf.page = page
                if page > 1:  # Skip page number on cover page
                    pdf.set_y(-15)
                    pdf.set_font("Helvetica", "I", 8)
                    pdf.cell(0, 10, f'Page {page} of {page_count}', align='C')
            
            # Generate PDF
            pdf.output(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating PDF: {str(e)}", exc_info=True)
            raise
    
    async def _export_to_docx(self, filename: str, messages: List[Dict]) -> str:
        """Export chat to DOCX"""
        output_path = os.path.join(self.export_dir, f"{filename}.docx")
        
        doc = docx.Document()
        
        # Set style
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        for msg in messages:
            role_icon = "👤" if msg["role"] == "user" else "🤖"
            timestamp = datetime.fromisoformat(msg["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
            
            # Add header
            header = doc.add_paragraph(f"{role_icon} {msg['role'].upper()} - {timestamp}")
            header.style = doc.styles['Heading 2']
            
            # Add content
            doc.add_paragraph(msg['content'])
            
            # Add separator
            doc.add_paragraph("_" * 80)
            doc.add_paragraph()
        
        doc.save(output_path)
        return output_path
    
    async def _export_to_txt(self, filename: str, messages: List[Dict]) -> str:
        """Export chat to TXT"""
        output_path = os.path.join(self.export_dir, f"{filename}.txt")
        
        formatted_text = self._format_messages(messages, 'txt')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(formatted_text)
            
        return output_path
    
    async def _export_to_markdown(self, filename: str, messages: List[Dict]) -> str:
        """Export chat to Markdown"""
        output_path = os.path.join(self.export_dir, f"{filename}.md")
        
        formatted_text = self._format_messages(messages, 'md')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(formatted_text)
            
        return output_path